from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Style helpers ─────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x16, 0x21, 0x3E)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x0F, 0x3C, 0x78)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11) if p.runs else None
    return p

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F0F0F0')
    p._p.pPr.append(shading)
    return p

def rule(label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    return p

def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'D0D8E8')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for run in cells[ci].paragraphs[0].runs:
                run.font.size = Pt(10)
    return t

def page_break():
    doc.add_page_break()

def divider():
    p = doc.add_paragraph('─' * 72)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('DEN SEJE APP')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('Avatar Creator — Production Spec Bundle v1')
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = meta.add_run('Visual Production Spec  ·  Art Bible  ·  Asset Production Pipeline')
r3.font.size = Pt(12)
r3.italic = True
r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = date_p.add_run('27. april 2026')
r4.font.size = Pt(11)
r4.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PART 1 — VISUAL PRODUCTION SPEC
# ══════════════════════════════════════════════════════════════════════════════
h1('DEL 1 — Visual Production Spec & Art Bible')
body('Gælder for alle assets i v1. Definerer proportioner, stil, farver, materiale og designregler for karakteren og de 4 core slots: hat, trøje, sko, inventory.')

# ── 1. Character Silhouette ───────────────────────────────────────────────────
h2('1. Character Silhouette Rules')

h3('Proportioner')
table(
    ['Mål', 'Regel'],
    [
        ['Head:body ratio', '1:6 — let oversized ift. realistisk 1:7.5, men IKKE chibi (1:3-4)'],
        ['Skulderbredde', '1.4× hoftbredde — svagt heroisk, ikke Fortnite-overdrevet'],
        ['Lemmer', '10% kortere end realistisk, rund tværsnitsform (ikke cylindrisk)'],
        ['Hænder', '115% af realistisk størrelse — bedre læsbarhed på lille skærm'],
        ['Fødder', '120% af realistisk størrelse — silhouette-forankring'],
        ['Halsvidde', 'Moderat — ALDRIG lang og tynd (Fortnite-markør)'],
    ]
)

h3('Stylization level')
body('Target: "Semi-stylized realism" — forenklede former, rene kanter, 15–20% ekstraggeration på nøglepunkter.')
body('Mere end Roblox (2/10), mindre end Fortnite (7/10). Target: 4–5/10. Former er organiske, ikke geometriske. Silhouette læser klart fra 50px thumbnail.')

h3('Hvad karakteren skal FØLES som visuelt')
for b in [
    'Selvsikker kropsposition — let bred stance, ikke defensiv',
    'Ren og optimistisk — ingen mørke undertoner i base design',
    'Skalerer godt til 2D profilbillede OG 3D roteret preview',
    'Kan bære alle slot-items uden at se "klædt udover" ud',
]:
    bullet(b)

h3('Hvad vi aktivt IKKE må ligne')
table(
    ['Forbud', 'Konkret regel'],
    [
        ['Roblox', 'Ingen blokke, ingen rektangulære lemmer, ingen plastik-cylinder-torso'],
        ['Fortnite', 'Ingen lang hals + brede skuldre + smal talje. Ingen den specifikke hoved-ægform'],
        ['Minecraft', 'Nul voxel-estetik. Nul firkantede hoved-proportiioner'],
        ['Chibi-style', 'Head:body må IKKE gå under 1:4.5'],
        ['Realistic human', 'Ingen fotorealistiske porer, ingen uncanny valley proportioner'],
    ]
)

# ── 2. Material & Rendering ───────────────────────────────────────────────────
h2('2. Material & Rendering Rules')

h3('Shading style')
body('Valg: Toon-adjacent stylized shading — 2–3 lysbånd, bløde skygger, ingen hård phong specular.')
body('Ikke: Fuld PBR (for tung, for realistisk, Fortnite-estetik). Ikke: Rå flat-color toon (for Roblox-adjacent).')
body('Konkret: Diffuse base + baked ambient occlusion + single rim light. Ingen realtime reflections på clothing i v1.')

h3('Texture style')
table(
    ['Parameter', 'Regel'],
    [
        ['Look', 'Håndmalet gradient-stil, IKKE scanned/fotograferet'],
        ['Støj', 'Subtil noise (max 5% opacity overlay) for at undgå plastik-look'],
        ['Detaljegrad', 'Skal læse klart ved 50×50px thumbnail'],
        ['Skygger', 'Bages ind — ingen realtime soft shadows på clothing i v1'],
        ['Specular', 'Kun på metal/plastik elementer. Aldrig på stof'],
    ]
)

h3('Mobile/browser performance')
table(
    ['Budget', 'Regel'],
    [
        ['Texture filtering', 'Bilinear (ikke trilinear) — performance'],
        ['Mip maps', 'Påkrævet for alle texturer ≥128px'],
        ['Draw calls pr. karakter', 'Max 4 i v1 (brug texture atlas)'],
        ['Shader complexity', 'Single pass, no multipass effects i v1'],
        ['Shadow casting', 'Karakter caster én simpel blob shadow'],
    ]
)

# ── 3. Color Language ─────────────────────────────────────────────────────────
h2('3. Color Language')

h3('Clothing primary palette')
table(
    ['Farve', 'Hex', 'Brug'],
    [
        ['Neutral White', '#F8F8F8', 'Base/accent — altid safe'],
        ['Charcoal', '#2C2C2C', 'Base — safe sort'],
        ['Slate Navy', '#1E3A5F', 'Base — skole-vibe'],
        ['Warm Grey', '#8A8A8A', 'Base/neutral'],
        ['Forest Green', '#2D6A4F', 'Base — naturlig'],
        ['Sky Blue', '#5BA4CF', 'Accent — pop'],
        ['Coral Red', '#E85D4A', 'Accent — pop'],
        ['Warm Yellow', '#F4C430', 'Accent — achievement'],
        ['Soft Lilac', '#9B8EC4', 'Accent'],
        ['Gold', '#C9A84C', 'Kun trofæer, kroner, achievement-items'],
    ]
)

h3('Forbudte farve-kombinationer')
table(
    ['Kombination', 'Årsag'],
    [
        ['Ren rød + hvid', 'Nike, Adidas, Coca-Cola'],
        ['Sort + guld', 'Jordan Brand'],
        ['Rød + sort + hvid (3-farve)', 'Stærk brand-association'],
        ['Grøn + gul', 'Brazil-flag, John Deere'],
        ['Orange + sort', 'Harley-Davidson feel'],
        ['Blå + orange + hvid', 'Mange sportshold'],
    ]
)

body('Regel: En enkelt farve er aldrig et problem. Det er KOMBINATIONER der skaber brand-association.')

# ── 4. Clothing Design Rules ──────────────────────────────────────────────────
h2('4. Clothing Design Rules')

h3('Hoodie/sweatshirt — tilladt')
for b in ['Kenguru-lomme (simpel)', 'Snøre-detalje (uden logo)', 'Ribkant ved manchetter og bund', 'Enkeltfarvet eller to-tone (body vs. hætte)', 'Subtil sømmønster som detalje']:
    bullet(b)

h3('Hoodie/sweatshirt — forbudt')
for b in ['Logo-placering på brystet (den klassiske brand-position)', 'Stor print på ryggen', 'Brand-specifik hætte-form', '"Box" tekst-placement anywhere']:
    bullet(b)

h3('Sneaker forbudte former')
table(
    ['Form', 'Brand det ligner'],
    [
        ['Swoosh-kurve nogen sted', 'Nike'],
        ['Tre parallelle striber, side', 'Adidas'],
        ['Jumpman-silhouette (lille figur)', 'Jordan'],
        ['Yeezy foam-tekstur på sål', 'Adidas Yeezy'],
        ['Air bubble-vindue i sål', 'Nike Air'],
        ['Ankelbeskytter med wing-detail', 'Jordan 1'],
        ['Velcro strap på ankel', 'Nike Killshot'],
    ]
)

body('Safe approach: Design en ren, rundet sneaker-silhouette fra scratch. Intet logo-territory. Enkel form. Tyk sål. Ingen eksisterende model-specifikke detaljer.')

# ── 5. Inventory Design Rules ─────────────────────────────────────────────────
h2('5. Inventory Design Rules')

h3('Principper')
for b in ['Identificeres på 0.5 sek ved 30px visning', 'Klar positiv/neutral konnotation', 'Ikke misforstås som våben fra nogen vinkel', 'Størrelse: 50–70% af karakter-håndens størrelse']:
    bullet(b)

h3('Godkendt inventory v1')
table(
    ['Item', 'School-signal'],
    [
        ['Åben bog', 'Høj'],
        ['Kæmpeblyant', 'Høj'],
        ['Spilkontroller (originalt design)', 'Neutral'],
        ['Gaming headset (holdes)', 'Neutral'],
        ['Trofæ/pokal', 'Høj (achievement)'],
        ['Pensel (stor)', 'Høj'],
        ['Fodbold', 'Neutral'],
        ['Skateboard (holdes)', 'Neutral'],
    ]
)

h3('Forbudt inventory')
table(
    ['Item', 'Årsag'],
    [
        ['Pistol (enhver stil)', 'Absolut forbud — skole'],
        ['Kniv/sværd som primær item', 'Skole-kontekst'],
        ['Granat', 'Absolut forbud'],
        ['Alkohol', 'Skole-kontekst'],
        ['Vape/cigaret', 'Absolut forbud'],
        ['Energidrik med specifik farve/form', 'Brand-association (Monster, Red Bull)'],
        ['Telefon med specificeret brand', 'Apple-logo, Samsung-layout'],
    ]
)

# ── 6. UI Preview Rules ───────────────────────────────────────────────────────
h2('6. UI Preview Rules')

h3('Viewport')
table(
    ['Parameter', 'Regel'],
    [
        ['Baggrund', 'Neutral gradient: lys grå (#E8E8E8) → hvid (#FFFFFF), top-til-bund'],
        ['Belysning', '3-point: primær front-venstre 60%, rim bag-højre 25%, fill blød nedefra 15%'],
        ['Standard vinkel', '3/4 view: 15° mod højre, 5° nedad'],
        ['Fuld krop synlig', 'Altid: hoved til fødder, 80–90% af frame-højde'],
        ['Post-processing', 'Ingen i v1'],
    ]
)

h3('Rotation')
for b in ['Y-akse kun (ingen X-tilt)', '360° fri rotation', 'Auto-snap til default 3/4 view efter 2.5s inaktivitet', 'Drag-velocity: dæmpes gradvist (ikke instant stop)']:
    bullet(b)

h3('Equip feedback')
for b in ['Item "popper" ind: scale 0.9 → 1.05 → 1.0 over 200ms', 'Hvid rim-highlight på det nye item i 400ms', 'Slot-ikon i UI viser checkmark i 800ms']:
    bullet(b)

# ── 7. Technical Constraints ──────────────────────────────────────────────────
h2('7. Technical Constraints')

h3('Polygon budget')
table(
    ['Asset', 'Tris (max)'],
    [
        ['Base karakter (krop, ingen tøj)', '3,000'],
        ['Hat', '1,200'],
        ['Hoodie/jacket', '1,800'],
        ['Sko (begge)', '2,000 (1,000 pr. sko)'],
        ['Inventory item', '800'],
        ['TOTAL v1 karakter', '≤ 10,000 tris'],
    ]
)

h3('Texture budget')
table(
    ['Asset', 'Resolution', 'Format'],
    [
        ['Base karakter', '512×512', 'PNG'],
        ['Hat', '256×256', 'PNG'],
        ['Shirt', '256×256', 'PNG'],
        ['Shoes (deles begge)', '256×256', 'PNG'],
        ['Inventory item', '128×128 eller 256×256', 'PNG'],
        ['Atlas mål (v1.1)', '2048×2048', 'PNG, alle items'],
    ]
)

h3('Attachment points')
table(
    ['Bone navn', 'Bruges til'],
    [
        ['attach_head', 'Hat, hjelm'],
        ['attach_hand_R', 'Inventory item (højre hånd)'],
        ['attach_hand_L', 'Reserveret til v2'],
        ['attach_foot_L', 'Venstre sko'],
        ['attach_foot_R', 'Højre sko'],
        ['attach_chest', 'Reserveret: badge/pin (v2)'],
        ['attach_back', 'Reserveret: rygsæk (v2)'],
    ]
)

h3('Export format')
table(
    ['Type', 'Format'],
    [
        ['Source files', '.blend (Blender) eller .fbx'],
        ['Web/browser export', '.glb (binary glTF 2.0)'],
        ['Texturer i glb', 'Embedded PNG'],
        ['Koordinatsystem', 'Y-up, meter-enheder, 1 unit = 1 meter'],
        ['Character højde', '1.75 meters i bind pose'],
    ]
)

# ── 8. Explicit Forbidden References ─────────────────────────────────────────
h2('8. Explicit Forbidden References')

h3('Games / IP')
table(
    ['Reference', 'Hvad der specifikt forbyder'],
    [
        ['Fortnite', 'Ægformet hoved, lang hals, brede skuldre+smal talje, specifikke skin-designs'],
        ['Roblox', 'Blokke, cilindrisk torso, firkantede hænder, plastik-shiny look'],
        ['Minecraft', 'Voxel, blokke, pixel-art style'],
        ['Pokémon', 'Nogen som helst karakter'],
        ['Among Us', 'Visir-form, rundede krop'],
        ['Fall Guys', 'Bønne-form, farve-kombination'],
        ['Overwatch/Valorant/LoL/Apex', 'Specifikke karakter/hero-designs'],
        ['Mario', 'Hat-silhouette, overall-design'],
        ['Zelda', 'Elfenøre, grøn keglehat'],
    ]
)

h3('Brands')
table(
    ['Reference', 'Hvad der specifikt forbyder'],
    [
        ['Nike', 'Swoosh-kurve, Air-bubble i sål, "Just Do It" typografi'],
        ['Adidas', 'Tre parallelle striber, trekant-logo, Trefoil'],
        ['Jordan', 'Jumpman, Wing-logo, Jordan 1 ankelbeskytter'],
        ['Supreme', 'Box-logo format (rektangel med serif font)'],
        ['Gucci', 'GG-monogram, dobbeltring'],
        ['Louis Vuitton', 'LV-monogram, Damier-mønster'],
        ['Off-White', 'Quotation marks som designelement, diagonal striber'],
        ['Balenciaga', 'Triple S-sålen specifikt'],
    ]
)

h3('Film / TV / Comics')
table(
    ['Reference', 'Forbud'],
    [
        ['Marvel', 'Nogen som helst superhelt-specifik dragt'],
        ['DC', 'Batman-maske, Superman-S, Wonder Woman'],
        ['Disney', 'Mickey-ører, alle Disney-karakterer'],
        ['Star Wars', 'Stormtrooper-hjelm, Darth Vader-maske, lightsaber-form'],
        ['Harry Potter', 'Specifik rund-bril + lyn-kombination (enkeltvis OK)'],
        ['Naruto', 'Karakterspecifikke markings, headband-design'],
    ]
)

h3('Symboler')
table(
    ['Kategori', 'Regel'],
    [
        ['Nazisymbolik', 'Absolut forbud: hagekors, SS-runer, enhver variant'],
        ['Politiske partisymboler', 'Ingen'],
        ['Religiøse symboler', 'Undgå som primær design-element'],
        ['Gangsymboler', 'Ingen'],
        ['Hadefulde symboler', 'Ingen'],
    ]
)

# ── 9. Starter Assets ─────────────────────────────────────────────────────────
h2('9. De 12 Starter-Assets — Prioriteret rækkefølge')

h3('Runde 1 — Foundation')
table(
    ['#', 'Asset', 'Slot', 'Begrundelse'],
    [
        ['1', 'Base karakter (krop + base outfit)', 'System', 'BLOKKERER ALT. Ingen assets bygges uden godkendt base mesh'],
        ['2', 'Plain hoodie (navy)', 'Shirt', 'Første tøjtest. Validerer skinning-pipeline og UV'],
        ['3', 'Baseball cap (charcoal)', 'Hat', 'Første attach_head test. Enkel geometri'],
        ['4', 'Low-top sneakers (white)', 'Shoes', 'Første attach_foot test. Validerer spejl-UV strategi'],
    ]
)

h3('Runde 2 — Validering')
table(
    ['#', 'Asset', 'Slot', 'Begrundelse'],
    [
        ['5', 'Spilkontroller (grey, originalt design)', 'Inventory', 'Første attach_hand_R test. Validerer inventory-pipeline'],
        ['6', 'Beanie (coral)', 'Hat', 'Tester farvevariant + hat-type #2'],
        ['7', 'High-top sneakers (black)', 'Shoes', 'Tester silhouette-variation på samme attachment setup'],
        ['8', 'Åben bog', 'Inventory', 'Tester flat/organic inventory-form, school-context signal'],
    ]
)

h3('Runde 3 — Komplet v1')
table(
    ['#', 'Asset', 'Slot', 'Begrundelse'],
    [
        ['9', 'Zip-up hoodie (white)', 'Shirt', 'Tester shirt variation #2, validerer zip-detail'],
        ['10', 'Varsity jacket (navy+gold)', 'Shirt', 'Tester to-tone material pipeline'],
        ['11', 'Mortarboard / eksamenshue', 'Hat', 'On-brand item, distinctive silhouette, hat #3'],
        ['12', 'Trofæ/pokal (gold)', 'Inventory', 'Tester metal-material, achievement-signal'],
    ]
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PART 2 — ASSET PRODUCTION PIPELINE SPEC
# ══════════════════════════════════════════════════════════════════════════════
h1('DEL 2 — Asset Production Pipeline Spec v1')
body('Definerer det komplette workflow fra idé til live asset. Gælder for alle 12 v1 assets.')

# ── 1. Asset Lifecycle ────────────────────────────────────────────────────────
h2('1. Asset Lifecycle')

body('Et asset gennemgår 12 diskrete stages. Ingen stage springes over. Ingen asset starter næste stage før nuværende er godkendt.')
code_block('BACKLOG → CONCEPT → BLOCKOUT → MODEL → UV → TEXTURE\n       → ATTACH → EXPORT → INTEGRATION → QA → APPROVAL → LIVE')

stages = [
    ('Stage 0: BACKLOG', 'Asset er på godkendt whitelist.', [
        'Design brief skrevet og godkendt',
        'Slot assignment bekræftet',
        'Copyright pre-screen gennemført',
        'Asset tilføjet til asset tracker',
    ]),
    ('Stage 1: CONCEPT', 'Design brief godkendt.', [
        'Minimum 3 ortografiske views (front, side, bag)',
        'Én 3/4-view',
        'Tegnet fra scratch — ingen trace af branded referencer',
        'Exit: Copyright Review Form godkendt + Blind test bestået',
    ]),
    ('Stage 2: BLOCKOUT', 'Concept Gate bestået.', [
        'Simpel boks-model, ingen detaljer',
        'Placeret mod base character i correct scale',
        'Screenshot ved 50×50px',
        'Exit: Silhouette læsbar, 2/2 reviewers godkender',
    ]),
    ('Stage 3: MODEL', 'Blockout godkendt.', [
        'Færdig geometri inden for polygon budget',
        'Ren quad-topologi (tris kun på export)',
        'Nul ngons',
        'Scale og rotation applied',
        'Exit: Geometry Gate bestået',
    ]),
    ('Stage 4: UV UNWRAP', 'Model godkendt.', [
        'UV layout i 0-1 space',
        'Korrekt island-padding',
        'Checker texture test screenshottede og vedlagt',
        'Exit: Ingen overlappende islands, ingen stretch over 20%',
    ]),
    ('Stage 5: TEXTURE', 'UV godkendt.', [
        'Tekstur i godkendt stil',
        '50×50px thumbnail PNG',
        'Udfyldt copyright visual review',
        'Exit: Inden for budget, thumbnail-test bestået, ingen brand-identifiable elementer',
    ]),
    ('Stage 6: RIG / ATTACH', 'Texture godkendt.', [
        'Clothing: skinnet til base rig',
        'Hat/sko/inventory: parented til korrekt attachment bone',
        'Rotation-test screenshots: 0°, 90°, 180°, 270°',
        'Exit: Attachment korrekt, nul clipping ved alle 4 rotation stops',
    ]),
    ('Stage 7: EXPORT', 'Rig/attach godkendt.', [
        '.glb fil (binary glTF 2.0, embedded textures)',
        'metadata.json udfyldt',
        'Filnavn overholder naming convention præcist',
        'Exit: Valid .glb (parser-test), metadata valid JSON',
    ]),
    ('Stage 8: INTEGRATION', 'Export godkendt af tech.', [
        'Asset loades i browser-preview',
        'Attachment bone valideres programmatisk',
        'Draw call count tælles',
        'Exit: Ingen console errors, draw calls ≤4',
    ]),
    ('Stage 9: QA', 'Integration godkendt.', [
        'Separat QA-person (ikke den artist der lavede asset)',
        'Fuld QA-tjekliste udfyldes',
        'Exit: Alle punkter PASS',
    ]),
    ('Stage 10: APPROVAL', 'QA bestået.', [
        'Tredelt godkendelse kræves',
        'Artist sign-off + Tech sign-off + Product sign-off',
        'Exit: Alle tre signaturer til stede',
    ]),
    ('Stage 11: STAGING', 'Approval komplet.', [
        'Deployed til staging environment',
        'Final visuel review i real product-kontekst',
        'Exit: Asset ser korrekt ud i fuld karakter-UI',
    ]),
    ('Stage 12: LIVE', 'Staging godkendt.', [
        'manifest.json opdateres: status "live"',
        'Deploy til produktion',
        'Monitoring-periode: 48 timer',
    ]),
]

for stage_name, entrance, bullets in stages:
    h3(stage_name)
    rule('Entrance criteria', entrance)
    for b in bullets:
        bullet(b)

# ── 2. Review Gates ───────────────────────────────────────────────────────────
h2('2. Review Gates')
body('Fire hårde gates. Et asset stopper ved enhver gate til den er bestået. Ingen fremgang.')

h3('Gate 1 — CONCEPT GATE')
body('Tidspunkt: Efter concept sketch, FØR nogen åbner modelleringsværktøj.')
table(
    ['Check', 'Regel'],
    [
        ['Copyright Review Form', 'Udfyldt, nul hits på blacklist'],
        ['Blind test', '3 reviewers: nul genkendelse af IP/brand'],
        ['Design brief sign-off', 'Product owner godkendt skriftligt'],
        ['Inspiration references', 'Listet og bekræftet som non-forbidden'],
    ]
)
body('Fail-konsekvens: Asset returneres til BACKLOG. Nyt concept kræves.')

h3('Gate 2 — GEOMETRY GATE')
body('Tidspunkt: Efter model + UV, FØR texture arbejde starter.')
table(
    ['Check', 'Regel'],
    [
        ['Polygon count', 'Under hård grænse'],
        ['Topology', 'Ingen ngons, ren quad-mesh'],
        ['UV compliance', 'Padding korrekt, ingen overlap'],
        ['Clipping test (rough)', 'Ingen åbenlys clipping mod base character'],
        ['Silhouette test', '50×50px — 2/2 reviewers godkender'],
    ]
)

h3('Gate 3 — ASSET COMPLETE GATE')
body('Tidspunkt: Efter texture + export, FØR engineering handoff.')
table(
    ['Check', 'Regel'],
    [
        ['Fuld QA self-check', 'Artist: alle punkter PASS'],
        ['Copyright visual review', 'Udfyldt, PASS'],
        ['Texture budget', 'Inden for grænse'],
        ['Thumbnail test', '50×50px: 2/2 reviewers'],
        ['Filnavn', 'Navnekonvention korrekt'],
        ['metadata.json', 'Valid, alle felter udfyldt'],
        ['.glb validering', 'Parser test bestået'],
    ]
)

h3('Gate 4 — RELEASE GATE')
body('Tidspunkt: Efter integration + QA, FØR staging/live.')
table(
    ['Check', 'Regel'],
    [
        ['Fuld QA checklist', 'Ekstern QA: alle punkter PASS'],
        ['Performance test', 'Load <200ms, draw calls ≤4'],
        ['Multi-browser test', 'Chrome, Firefox, Safari — alle pass'],
        ['Tredelt approval', 'Alle tre signaturer til stede'],
    ]
)

# ── 3. Acceptance Criteria ────────────────────────────────────────────────────
h2('3. Acceptance Criteria')
body('Alle regler er objektive og målbare. PASS/FAIL — ingen "næsten OK".')

h3('Polygon Budget')
table(
    ['Slot', 'Max tris', 'Konsekvens'],
    [
        ['Hat', '1,200', 'HARD FAIL — returneres til MODEL'],
        ['Shirt/jacket', '1,800', 'HARD FAIL'],
        ['Shoes (begge samlet)', '2,000 (1,000 pr. sko)', 'HARD FAIL'],
        ['Inventory', '800', 'HARD FAIL'],
        ['Base character', '3,000', 'HARD FAIL'],
        ['Total fuld karakter', '10,000', 'HARD FAIL'],
    ]
)

h3('Texture Budget')
table(
    ['Asset', 'Max resolution', 'Max filstørrelse'],
    [
        ['Base character', '512×512', '512KB'],
        ['Hat', '256×256', '128KB'],
        ['Shirt', '256×256', '128KB'],
        ['Shoes (deles)', '256×256', '128KB'],
        ['Inventory', '256×256', '128KB'],
        ['Total karakter', '—', '≤1MB'],
    ]
)

h3('Øvrige acceptance criteria')
table(
    ['Test', 'Pass-krav'],
    [
        ['Thumbnail readability', '2/2 reviewers identificerer korrekt inden for 3 sekunder'],
        ['Clipping test', 'Ingen clipping >2px ved nogen af 24 rotationsstops (0°–360° i 15° inkrementer)'],
        ['Copyright review', '3/3 blind test reviewers ser intet specifikt brand/IP/karakter'],
        ['Rotation test', 'Asset komplet fra alle vinkler, ingen hul bagside'],
        ['Performance — load', '<100ms isoleret, <300ms fuld karakter (throttled mobil)'],
        ['Performance — FPS', 'Stabil 60fps Chrome, ≥30fps Safari ved rotation'],
        ['Attachment', 'Hat ≤2px fra hoved. Shoes ≤0.01 units fra ground. Inventory ingen clip'],
    ]
)

# ── 4. Artist Workflow ────────────────────────────────────────────────────────
h2('4. Artist Workflow')

steps = [
    ('Step 1 — Modtag og læs design brief',
     ['Læs alle felter i design brief', 'Verificer slot og naming convention er forstået',
      'Læs "NOT allowed to look like" listen', 'Stil spørgsmål NU — ikke efter modeling starter']),
    ('Step 2 — Udfyld Copyright Review Form del 1',
     ['List ALLE referencer du vil lade dig inspirere af', 'Kør dem mod blacklisten selv',
      'Submit for Concept Gate review', 'VENT på godkendelse. Start ikke modeling.']),
    ('Step 3 — Concept sketch',
     ['Tegn fra scratch. Ingen trace.', 'Minimum: front, side, bag, 3/4',
      'Indikér colorway', 'Submit til Concept Gate']),
    ('Step 4 — Vent på Gate 1 godkendelse',
     ['Godkendt? → Fortsæt', 'Afvist? → Tilbage til step 2 med ændret concept']),
    ('Step 5 — Blockout i Blender',
     ['Åbn base_character_v1.0.blend (ALDRIG modificér denne fil)',
      'Link base character som reference object (ikke merge)',
      'Byg simpelt blockout i korrekt skala',
      'Screenshot: 50×50px silhouette test',
      'Send screenshot til team']),
    ('Step 6 — Final model',
     ['Byg ud fra blockout, hold poly-budget for øje løbende',
      'Check poly count kontinuerligt (Blender: N-panel → Statistics)',
      'Ingen ngons — verificer med Mesh > Cleanup > Fill Holes',
      'Apply scale og rotation: Object → Apply → All Transforms',
      'Geometry Gate submission: poly count screenshot + .blend']),
    ('Step 7 — UV Unwrap',
     ['Unwrap med Smart UV Project som base, manuel finpudsning',
      'Pack islands, verificer minimum padding',
      'Checker texture test: åbenlys stretch = fix det',
      'Screenshot af UV layout og checker test — vedlæg til Geometry Gate']),
    ('Step 8 — Texture',
     ['Mal ved 512px, eksportér til 256px (downscale EFTER maling)',
      'Ingen brug af fotografiske sources',
      'Stil: stylized gradient, håndmalet look',
      'Udfyld Copyright Review Form del 2 (visuel review)',
      'Export 50×50px thumbnail PNG',
      'Submit til Asset Complete Gate review']),
    ('Step 9 — Rig / Attach',
     ['Hat/sko/inventory (parented): Parent asset til korrekt attach bone. Verify position. Roter 0°/90°/180°/270° og screenshot hvert stop.',
      'Shirt (skinned): Skin mesh til base rig bones. Verify skinning: ingen stiff polygons.',
      'Alle 4 rotation screenshots vedlægges submission']),
    ('Step 10 — Self-QA',
     ['Kør fuld QA-tjekliste på dit eget asset',
      'Alle punkter skal markeres PASS',
      'Er der et FAIL: fix det FØR du submitter',
      'Submit aldrig et asset med kendte fejl']),
    ('Step 11 — Export',
     ['File > Export > glTF 2.0 (.glb)',
      'Settings: Format = glTF Binary, Include = Selected Objects, Apply Modifiers = ON, Textures embed = ON',
      'Filnavn: exakt efter naming convention',
      'Udfyld metadata.json',
      'Submit begge filer til engineering']),
]

for step_title, step_bullets in steps:
    h3(step_title)
    for b in step_bullets:
        bullet(b)

# ── 5. Engineering Handoff ────────────────────────────────────────────────────
h2('5. Engineering Handoff')

h3('Handoff-pakke per asset')
code_block(
    'hat_baseballcap_navy_v1.0.glb\n'
    'hat_baseballcap_navy_v1.0_metadata.json\n'
    'hat_baseballcap_navy_v1.0_qa.md\n'
    'hat_baseballcap_navy_v1.0_copyright.md'
)

h3('metadata.json format')
code_block(
    '{\n'
    '  "asset_id": "hat_baseballcap_navy",\n'
    '  "slot": "hat",\n'
    '  "display_name_da": "Baseball kasket",\n'
    '  "color_variant": "navy",\n'
    '  "poly_count": 856,\n'
    '  "texture_resolution": "256x256",\n'
    '  "texture_size_kb": 48,\n'
    '  "attachment_bone": "attach_head",\n'
    '  "school_safe": true,\n'
    '  "copyright_reviewed": true,\n'
    '  "version": "1.0.0",\n'
    '  "status": "pending_integration"\n'
    '}'
)

h3('Engineer integration checklist')
for b in [
    '.glb parser test: file parses without error',
    'Attachment bone present in scene: attach_{slot}',
    'Texture resolution matches metadata',
    'Poly count matches metadata (verify via glTF validator)',
    'Asset loads in browser: zero console errors',
    'Asset loads in <100ms (isolated, throttled)',
    'Attachment position visually correct on base character',
    'manifest.json updated: status "staging"',
    'Staging deploy completed',
]:
    bullet(b)

# ── 6. Naming Conventions ─────────────────────────────────────────────────────
h2('6. Naming Conventions')

h3('Filnavn — production assets')
code_block('{slot}_{assetname}_{colorvariant}_v{major}.{minor}.{ext}\n\nEksempler:\nhat_baseballcap_navy_v1.0.glb\nshirt_hoodie_coral_v1.0.glb\nshoe_lowtop_white_v1.0.glb\ninventory_controller_grey_v1.0.glb')

h3('Mappestruktur')
code_block(
    'assets/\n'
    '└── avatar/\n'
    '    ├── manifest.json\n'
    '    ├── v1/\n'
    '    │   ├── hat/\n'
    '    │   ├── shirt/\n'
    '    │   ├── shoe/\n'
    '    │   └── inventory/\n'
    '    ├── source/\n'
    '    │   ├── hat/\n'
    '    │   ├── shirt/\n'
    '    │   ├── shoe/\n'
    '    │   └── inventory/\n'
    '    ├── qa/\n'
    '    ├── concept/\n'
    '    └── base/\n'
    '        ├── base_character_v1.0.blend  ← LOCKED, read-only\n'
    '        ├── base_character_v1.0.glb\n'
    '        └── character_rig_v1.0.blend   ← LOCKED, read-only'
)

h3('Navneregler — absolutte')
for b in [
    'Kun lowercase',
    'Kun underscore som separator — aldrig bindestreger, aldrig mellemrum',
    'Ingen special characters @#$%^&*()',
    'Slot-navne: hat, shirt, shoe, inventory — aldrig pluralis',
    'Farvenavne: beskrivende ord — aldrig hex-koder i filnavn',
    'Versionsnumre: altid v{major}.{minor} — v1.0, v1.1, v2.0',
    'Ingen "final", "final2", "finalFINAL" — brug versionstal',
]:
    bullet(b)

h3('Godkendte farvevarianter (v1)')
code_block('white, black, charcoal, grey, navy, forest, coral, sky, lilac, gold, silver, warm_yellow')

# ── 7. Version Control Strategy ───────────────────────────────────────────────
h2('7. Version Control Strategy')

h3('To repositories')
body('Asset repo (binære filer — Git LFS): .glb, .blend, .png, .psd filer. Én branch per asset. Merge til main KUN efter fuld Gate 4 approval. Aldrig commit WIP .glb til main.')
body('Code repo (JSON, configs, frontend): manifest.json og alle metadata.json filer. Opdateres kun via PR.')

h3('Branch strategi')
code_block(
    'main\n'
    '  (kun godkendte, live assets)\n'
    '  ↑\n'
    'release/v1\n'
    '  (staging branch)\n'
    '  ↑\n'
    'asset/hat-baseballcap-navy\n'
    'asset/shirt-hoodie-coral\n'
    'asset/shoe-lowtop-white\n'
    'asset/inventory-controller-grey'
)

h3('Semantic versioning')
table(
    ['Version', 'Hvornår'],
    [
        ['v1.0', 'Første godkendte version'],
        ['v1.1', 'Minor fix: tekstur-tweak, lille geometri-fix'],
        ['v1.2', 'Endnu et minor fix'],
        ['v2.0', 'Redesign: ny shape, ny topologi, visuelt anderledes'],
    ]
)
body('Gamle versioner slettes aldrig. De sættes til "status": "deprecated" i manifest.')

# ── 8. Copyright Review Workflow ──────────────────────────────────────────────
h2('8. Copyright Review Workflow')
body('To obligatoriske reviews per asset. Ingen single-review undtagelser.')

h3('Review 1 — Concept Stage (Gate 1)')
body('Hvem: Artist (initialt) + Product (final godkendelse).')
for b in [
    'List alle referencer brugt til inspiration',
    'Gennemgå forbidden reference blacklist punkt for punkt',
    'Check alle farve-kombinationer mod brand-color liste',
    'Pass: nul hits på blacklist',
    'Fail: ét enkelt hit = FAIL, uanset hvor lille',
]:
    bullet(b)

h3('Review 2 — Completed Asset (Gate 3) — BLIND TEST')
body('Hvem: Tre separate reviewers. Ingen kontekst gives.')
for b in [
    'Send rendered asset preview — ingenting andet, ingen filnavn, ingen context',
    'Spørg: "Hvad minder dette dig om? Ligner det noget specifikt?"',
    'Notér alle spontane associationer',
    'Pass: 3/3 reviewers nævner intet specifikt brand, IP eller karakter',
    'Fail: Én reviewer nævner en specifik reference → asset returneres til concept',
]:
    bullet(b)

h3('Hvad en artist ALDRIG må gøre')
for b in [
    'Bruge brand-specifikke fotos som direkte reference/trace',
    '"Inspirere sig" af screenshots fra forbudte spil direkte',
    'Bruge AI-genererede billeder som reference hvis prompten inkluderede forbudte referencer',
    'Lade sig presse til at godkende noget der "ligner lidt"',
]:
    bullet(b)

# ── 9. QA Checklist ───────────────────────────────────────────────────────────
h2('9. QA Checklist')
body('Udfyldes af QA-person (ikke artist). Alle punkter kræver eksplicit PASS eller FAIL + noter.')

qa_sections = [
    ('GEOMETRY', [
        'Polygon count under hård grænse — faktisk count: ___ / Max: ___',
        'Ingen ngons i final mesh',
        'Ingen floating geometry',
        'Ingen duplicate vertices',
        'Alle transforms applied (scale=1, rot=0)',
        'Origin point korrekt',
    ]),
    ('UV', [
        'Alle UVs i 0-1 space',
        'Minimum 4px padding ved 256px',
        'Ingen overlappende UV islands (undtagen godkendte spejle)',
        'UV stretch under 20% (checker test)',
    ]),
    ('TEXTURE', [
        'Filstørrelse under budget: ___ KB / Max: ___',
        'Resolution er power-of-2',
        'Thumbnail 50×50px: reviewer 1 identificerer korrekt',
        'Thumbnail 50×50px: reviewer 2 identificerer korrekt',
        'Ingen fotografiske texture sources',
        'Style-compliance: stylized, hand-painted look',
    ]),
    ('TEKNISK', [
        '.glb parser test bestået (glTF Validator)',
        'Textures embedded i .glb',
        'Attachment bone til stede, korrekt navn',
        'Loader i browser: nul console errors',
        'Load time isoleret: ___ ms / Max: 100ms',
    ]),
    ('VISUEL', [
        'Clipping test 0°: ingen clipping',
        'Clipping test 90°: ingen clipping',
        'Clipping test 180°: ingen clipping',
        'Clipping test 270°: ingen clipping',
        'Scale korrekt: ingen float, ingen sink',
        'Korrekt ved full-body zoom',
        'Korrekt ved detail-zoom',
        'Asset er ikke hul bagpå',
    ]),
    ('COPYRIGHT', [
        'Concept-stage copyright form: PASS på fil',
        'Final blind test: 3/3 reviewers: ingen IP/brand',
        'Nul blacklist-hits bekræftet',
        'Colorway-kombination godkendt',
    ]),
    ('SCHOOL-SAFE', [
        'Ingen aggressiv/voldelig symbolik',
        'Ingen politisk symbolik',
        'Ingen religiøs symbolik som primær design',
        'Ingen brandlogoer',
        'School-safe standard opfyldt',
    ]),
    ('METADATA', [
        'metadata.json valid JSON',
        'Alle required felter udfyldt',
        'asset_id matcher filnavn præcist',
        'slot-felt korrekt',
    ]),
]

for section_title, checks in qa_sections:
    h3(section_title)
    for c in checks:
        bullet(f'[ ] PASS  [ ] FAIL   {c}')

# ── 10. Approval Ownership ────────────────────────────────────────────────────
h2('10. Approval Ownership')
body('Tre-niveau godkendelse. Alle tre kræves. Ingen kan erstatte en anden.')

table(
    ['Niveau', 'Hvem', 'Ansvar', 'Hvad de IKKE godkender'],
    [
        ['1 — Artist', 'Artist der lavede asset', 'Geometry, UV, texture kvalitet, style compliance, self-QA, export korrekt', 'Copyright, school-safe, release-beslutning'],
        ['2 — Tech', 'Frontend-ansvarlig engineer', 'Filformat gyldigt, performance specs, attachment points, integration test, manifest.json opdateret', 'Visual quality, copyright'],
        ['3 — Product', 'Product owner (Moeller888)', 'Final copyright, school-safe, visuel konsistens, release-beslutning', 'Technical implementation'],
    ]
)

body('Ingen asset går live uden alle tre signaturer. Hvis product owner er utilgængelig: asset venter. Der er ingen bypass.')

# ── 11. Rollback Strategy ─────────────────────────────────────────────────────
h2('11. Rollback Strategy')

h3('Severity klassificering')
table(
    ['Level', 'Årsag', 'Reaktionstid'],
    [
        ['S1', 'Copyright/legal bekræftet', 'Øjeblikkelig — under 1 time'],
        ['S2', 'School-safety issue', 'Under 2 timer'],
        ['S3', 'Visuel fejl, broken experience', 'Under 24 timer'],
        ['S4', 'Minor visual issue', 'Næste planlagte release'],
    ]
)

h3('Rollback procedure (S1 og S2)')
for b in [
    'manifest.json: asset status "live" → "rollback" (øjeblikkelig deploy)',
    'System viser: ingen item i den slot (slot er tom men karakter fungerer)',
    'Alternativt: fallback til forrige version hvis den eksisterer',
    'Ingen elev efterlades med broken karakter-visning',
    'Intern incident rapport udfyldes inden for 24 timer',
    'Post-mortem: hvad fejlede i review-pipeline?',
]:
    bullet(b)

h3('Kritisk fallback-krav')
body('Base character MÅ se komplet og neutral ud med nul items equipped i alle slots. Dette er et produktionskrav — ikke et nice-to-have. Det sikrer at rollback aldrig bryder en elevs oplevelse.')

# ── 12. Mindste korrekte v1 implementation ────────────────────────────────────
h2('12. Mindste korrekte v1 implementation')
body('Det mindste professionelle system der er korrekt og fungerer for 1–3 person team:')
for b in [
    '1. Asset tracker — ét Google Sheet (asset_id | stage | gate_1 | gate_2 | gate_3 | gate_4 | owner | notes)',
    '2. Mappestruktur — som defineret i navnekonventioner (ingen special tools, bare mapper)',
    '3. Fire template-filer: copyright_review_template.md, qa_checklist_template.md, metadata_template.json, design_brief_template.md',
    '4. manifest.json — single source of truth for live assets',
    '5. Base character + rig — locked, read-only, alle arbejder fra disse',
]:
    bullet(b)

body('Ingen Jira. Ingen Notion. Ingen CI/CD pipelines i v1. De kan tilføjes i v2 når volumen kræver det.')

# ── 13. Den ene vigtigste første handling ─────────────────────────────────────
h2('13. Den ENE vigtigste første handling')

p = doc.add_paragraph()
r = p.add_run('Lås base character meshet. Intet andet sker før det.')
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0xC0, 0x20, 0x20)

for b in [
    '1. Base character mesh modelleres',
    '2. Attachment points defineres og navngives',
    '3. Base rig oprettes med alle 16 bones',
    '4. base_character_v1.0.blend og base_character_v1.0.glb gemmes',
    '5. Begge filer sættes til read-only',
    '6. Product sign-off: Moeller888 godkender base character',
    '7. Derefter og KUN derefter: første asset-brief oprettes',
]:
    bullet(b)

body('Årsag: Hvert eneste andet punkt i denne spec afhænger af en stabil base character. Polygon-budgetter kalibreres mod den. Clipping-tests kræver den. Attachment points er defineret på den. UV-referencer skaleres mod den. Rig-hierarkiet er den.')
body('Ændres base character efter assets er bygget: alle 12 assets skal retestes og potentielt reworkes. Lås den først.')

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = r'C:\Users\Bruger\Documents\DEN SEJE APP\DEN SEJE APP\docs\Avatar_Creator_Production_Spec_v1.docx'
os.makedirs(os.path.dirname(out_path), exist_ok=True)
doc.save(out_path)
print(f'Saved: {out_path}')
